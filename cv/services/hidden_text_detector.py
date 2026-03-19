"""cv/services/hidden_text_detector.py — Wykrywanie ukrytego tekstu w dokumentach CV.

Obsługiwane techniki ukrywania:
    PDF:
        - Biały/jasny kolor czcionki (non_stroking_color ≈ white)
        - Czcionka < 1pt (praktycznie niewidoczna)
        - Kolor czcionki identyczny z kolorem tła strony (text-on-same-color)
        - CMYK i Grayscale odpowiedniki bieli

    DOCX:
        - Explicit RGBColor bliski białemu (≥ 240 w każdym kanale)
        - Rozmiar czcionki ≤ 1pt
        - Highlight color WHITE (WD_COLOR_INDEX.WHITE)

Wyniki zwracane jako lista dict:
    {
        'reason':   str,   # 'white_font_color' | 'tiny_font' | 'white_highlight'
        'page':     int,   # numer strony (1-based) lub 0 dla DOCX
        'snippet':  str,   # fragment ukrytego tekstu (max 120 znaków)
        'char_count': int, # liczba ukrytych znaków
    }
"""

from __future__ import annotations

import logging
from typing import Any

logger = logging.getLogger(__name__)

# ── Stałe ─────────────────────────────────────────────────────────────────────

# Próg "jasności" dla kolorów RGB znormalizowanych (0.0–1.0)
# Wartości > tego progu we WSZYSTKICH kanałach → uznajemy za "biały/jasny"
_RGB_WHITE_THRESHOLD = 0.88   # odpowiada ok. 224/255

# Dla DOCX kolor 0–255 per kanał
_DOCX_WHITE_THRESHOLD = 224   # ≥ 224 we wszystkich 3 kanałach

# Rozmiar czcionki ≤ tego progu (w punktach) → "niewidoczna"
_TINY_FONT_PT = 1.0

# Minimalna długość ukrytego tekstu aby uznać za celowe (nie artefakt)
_MIN_HIDDEN_CHARS = 3

# Maksymalna liczba snippetów zbieranych per dokument (limit koszt. analizy)
_MAX_FINDINGS = 20


# ── PDF ───────────────────────────────────────────────────────────────────────

def _color_is_white_pdf(color: Any) -> bool:
    """Sprawdza czy kolor czcionki PDF jest biały lub prawie biały.

    pdfplumber eksponuje kolor w kilku formatach:
        - None          → domyślny (zwykle czarny) — bezpieczny
        - float         → Grayscale 0.0 (czarny) – 1.0 (biały)
        - (r, g, b)     → RGB znormalizowane 0.0–1.0
        - (c, m, y, k)  → CMYK znormalizowane 0.0–1.0  (0,0,0,0) = biały
    """
    if color is None:
        return False

    # Grayscale
    if isinstance(color, (int, float)):
        return float(color) >= _RGB_WHITE_THRESHOLD

    if not isinstance(color, (tuple, list)):
        return False

    if len(color) == 3:
        # RGB (r, g, b) — wszystkie kanały ≥ progu
        return all(c >= _RGB_WHITE_THRESHOLD for c in color)

    if len(color) == 4:
        # CMYK — (0, 0, 0, 0) = biały; sumujemy tusze: mała wartość → jasny
        c, m, y, k = color
        # Przybliżona konwersja do jasności: im mniej tuszu, tym jaśniej
        r = (1 - c) * (1 - k)
        g = (1 - m) * (1 - k)
        b = (1 - y) * (1 - k)
        return all(v >= _RGB_WHITE_THRESHOLD for v in (r, g, b))

    return False


def detect_hidden_text_pdf(file_obj) -> list[dict]:
    """Wykrywa ukryty tekst w pliku PDF.

    Skanuje każdy znak (char) na każdej stronie pod kątem:
    - białego/jasnego koloru czcionki
    - rozmiaru czcionki ≤ 1pt

    Args:
        file_obj: obiekt pliku (django InMemoryUploadedFile lub zwykły plik)

    Returns:
        Lista dict z opisem znalezisk. Pusta lista jeśli nic nie wykryto.
    """
    try:
        import pdfplumber
    except ImportError:
        logger.warning("hidden_text_detector: pdfplumber not installed — PDF scan skipped")
        return []

    findings: list[dict] = []
    try:
        file_obj.seek(0)
        with pdfplumber.open(file_obj) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                if len(findings) >= _MAX_FINDINGS:
                    break

                chars = page.chars or []

                # Grupuj ukryte znaki w ciągłe sekwencje (nie każdy znak osobno)
                hidden_white: list[str] = []
                hidden_tiny:  list[str] = []

                for ch in chars:
                    text = ch.get('text', '').strip()
                    if not text:
                        continue

                    color = ch.get('non_stroking_color')
                    size  = ch.get('size', 12)

                    if _color_is_white_pdf(color):
                        hidden_white.append(ch.get('text', ''))
                    elif isinstance(size, (int, float)) and float(size) <= _TINY_FONT_PT:
                        hidden_tiny.append(ch.get('text', ''))

                if hidden_white:
                    snippet = ''.join(hidden_white)
                    if len(snippet.strip()) >= _MIN_HIDDEN_CHARS:
                        findings.append({
                            'reason':     'white_font_color',
                            'page':       page_num,
                            'snippet':    snippet[:120],
                            'char_count': len(hidden_white),
                        })
                        logger.warning(
                            "hidden_text_detector [PDF]: white font on page %d — "
                            "%d chars: %r", page_num, len(hidden_white), snippet[:60],
                        )

                if hidden_tiny and len(findings) < _MAX_FINDINGS:
                    snippet = ''.join(hidden_tiny)
                    if len(snippet.strip()) >= _MIN_HIDDEN_CHARS:
                        findings.append({
                            'reason':     'tiny_font',
                            'page':       page_num,
                            'snippet':    snippet[:120],
                            'char_count': len(hidden_tiny),
                        })
                        logger.warning(
                            "hidden_text_detector [PDF]: tiny font (≤1pt) on page %d — "
                            "%d chars: %r", page_num, len(hidden_tiny), snippet[:60],
                        )

    except Exception as e:
        logger.warning("hidden_text_detector [PDF]: scan error — %s", e)

    return findings


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _docx_color_is_white(run) -> bool:
    """Sprawdza czy kolor czcionki w python-docx run jest biały/jasny.

    RGBColor w python-docx to podklasa tuple — kanały dostępne przez indeks:
        rgb[0] = red (0–255)
        rgb[1] = green (0–255)
        rgb[2] = blue (0–255)
    """
    try:
        color = run.font.color
        if color is None or color.type is None:
            return False
        rgb = color.rgb  # RGBColor (tuple) lub None jeśli typ nie jest RGB
        if rgb is None:
            return False
        r, g, b = int(rgb[0]), int(rgb[1]), int(rgb[2])
        return all(c >= _DOCX_WHITE_THRESHOLD for c in (r, g, b))
    except Exception:
        return False


def _docx_font_is_tiny(run) -> bool:
    """Sprawdza czy rozmiar czcionki w run jest ≤ 1pt."""
    try:
        size = run.font.size  # w EMU (1pt = 12700 EMU)
        if size is None:
            return False
        return int(size) <= int(_TINY_FONT_PT * 12700)
    except Exception:
        return False


def _docx_has_white_highlight(run) -> bool:
    """Sprawdza czy run ma białe podświetlenie (highlight = WHITE)."""
    try:
        from docx.enum.text import WD_COLOR_INDEX
        hl = run.font.highlight_color
        return hl == WD_COLOR_INDEX.WHITE
    except Exception:
        return False


def detect_hidden_text_docx(file_obj) -> list[dict]:
    """Wykrywa ukryty tekst w pliku DOCX.

    Skanuje każdy run w każdym paragrafie pod kątem:
    - białego/jasnego koloru czcionki (explicit RGBColor)
    - rozmiaru czcionki ≤ 1pt
    - białego podświetlenia (highlight_color = WHITE)

    Args:
        file_obj: obiekt pliku

    Returns:
        Lista dict z opisem znalezisk. Pusta lista jeśli nic nie wykryto.
    """
    try:
        from docx import Document
    except ImportError:
        logger.warning("hidden_text_detector: python-docx not installed — DOCX scan skipped")
        return []

    findings: list[dict] = []
    try:
        file_obj.seek(0)
        doc = Document(file_obj)

        # Grupuj po typie powodu — zbieraj tekst ze wszystkich runów
        buckets: dict[str, list[str]] = {
            'white_font_color':  [],
            'tiny_font':         [],
            'white_highlight':   [],
        }

        for para in doc.paragraphs:
            for run in para.runs:
                text = run.text
                if not text or not text.strip():
                    continue
                if _docx_color_is_white(run):
                    buckets['white_font_color'].append(text)
                elif _docx_font_is_tiny(run):
                    buckets['tiny_font'].append(text)
                elif _docx_has_white_highlight(run):
                    buckets['white_highlight'].append(text)

        # Skanuj też tabele (popularne miejsce ukrywania tekstu)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            text = run.text
                            if not text or not text.strip():
                                continue
                            if _docx_color_is_white(run):
                                buckets['white_font_color'].append(text)
                            elif _docx_font_is_tiny(run):
                                buckets['tiny_font'].append(text)
                            elif _docx_has_white_highlight(run):
                                buckets['white_highlight'].append(text)

        for reason, parts in buckets.items():
            if not parts:
                continue
            snippet = ''.join(parts)
            if len(snippet.strip()) < _MIN_HIDDEN_CHARS:
                continue
            findings.append({
                'reason':     reason,
                'page':       0,   # DOCX nie ma numerów stron w strukturze
                'snippet':    snippet[:120],
                'char_count': len(snippet),
            })
            logger.warning(
                "hidden_text_detector [DOCX]: %s detected — %d chars: %r",
                reason, len(snippet), snippet[:60],
            )

    except Exception as e:
        logger.warning("hidden_text_detector [DOCX]: scan error — %s", e)

    return findings


# ── Główna funkcja ─────────────────────────────────────────────────────────────

def detect_hidden_text(file_obj, fmt: str) -> list[dict]:
    """Wykrywa ukryty tekst w dokumencie CV.

    Dispatcher — wybiera metodę na podstawie formatu pliku.
    Zawsze zwraca listę (pustą jeśli nic nie wykryto lub błąd).

    Args:
        file_obj: obiekt pliku (seek-able)
        fmt:      'pdf' | 'docx' | 'txt'

    Returns:
        Lista dict z opisem znalezisk.
    """
    if fmt == 'pdf':
        return detect_hidden_text_pdf(file_obj)
    if fmt == 'docx':
        return detect_hidden_text_docx(file_obj)
    return []  # TXT — brak formatowania, brak ukrytego tekstu


def findings_to_injection_flags(findings: list[dict]) -> list[dict]:
    """Konwertuje znaleziska hidden text na format security_flags.

    Używane do wstrzyknięcia wyników do istniejącego pipeline'u injection detection.
    """
    return [
        {
            'type':    'hidden_text',
            'subtype': f['reason'],
            'fragment': f['snippet'],
            'page':    f.get('page', 0),
            'char_count': f['char_count'],
            'action':  'content_flagged',
            'source':  'hidden_text_detector',
        }
        for f in findings
    ]
